import json, io, os
import hashlib
from datetime import datetime, timedelta
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlmodel import Session, select
from models import Case, TimelineEvent, AuditLog
from database import get_session
from routers.auth import get_current_user, _audit
from models import User
from ai_router import call_ai, call_ai_json

router = APIRouter(prefix="/api/reports", tags=["reports"])


def severity_color_rgb(severity: str):
    return {
        "critical": (192, 57, 43),
        "high":     (234, 179, 8),
        "medium":   (167, 139, 250),
        "low":      (34, 197, 94),
        "info":     (113, 113, 122),
    }.get(severity, (113, 113, 122))


def ai_enrich_case(case: Case) -> dict:
    """
    Use AI to generate detailed narrative sections if not already present.
    Returns a dict with enriched text fields.
    """
    iocs  = json.loads(case.iocs or "[]")
    mitre = json.loads(case.mitre_techniques or "[]")

    # If we already have AI summaries, just enhance them for the report
    existing_exec = case.ai_executive_summary or ""
    existing_tech = case.ai_technical_summary or ""

    prompt = f"""You are a senior SOC analyst writing a formal incident report.
Generate detailed, professional narrative sections for the following case.

Case: {case.case_number} — {case.title}
Severity: {case.severity.upper()} | Type: {case.incident_type}
Affected Systems: {case.affected_systems}
Analyst: {case.analyst_name} | Customer: {case.customer_name}
Description: {case.description[:600]}
Findings: {case.findings[:800]}
Commands Run: {(case.commands_run or '')[:600]}
IOCs ({len(iocs)}): {json.dumps(iocs[:8])}
MITRE Techniques: {json.dumps(mitre)}
Existing Summary: {existing_exec[:300]}

Generate a comprehensive, well-written incident report. Respond ONLY with valid JSON:
{{
  "executive_summary": "3-4 professional sentences for senior management and legal. Include business impact, attack vector, and outcome.",
  "technical_narrative": "4-6 detailed technical paragraphs covering: initial detection, attack analysis, threat actor TTPs, evidence collected, containment actions, and forensic findings.",
  "attack_timeline_summary": "2-3 sentences summarising the chronological attack progression.",
  "impact_assessment": "Paragraph assessing: data loss, financial impact, service disruption, reputational risk, regulatory implications.",
  "threat_actor_profile": "2-3 sentences on threat actor sophistication, likely motivation, and attribution confidence.",
  "remediation_summary": "3-4 specific, actionable remediation steps taken or recommended.",
  "lessons_learned": "2-3 specific operational improvements identified from this incident.",
  "risk_rating": "Overall risk rating with justification (Critical/High/Medium/Low and why)"
}}"""

    result = call_ai_json("report", prompt, temperature=0.3, max_tokens=2000)
    if result.get("parse_error"):
        return {
            "executive_summary": existing_exec or case.description,
            "technical_narrative": existing_tech or case.findings,
            "attack_timeline_summary": "",
            "impact_assessment": "",
            "threat_actor_profile": "",
            "remediation_summary": case.recommendations or "",
            "lessons_learned": "",
            "risk_rating": f"{case.severity.upper()} — Score: {case.ai_severity_score or 'N/A'}/100",
        }
    return result


@router.get("/{case_id}/pdf")
def download_pdf(case_id: int, session: Session = Depends(get_session),
                 user: User = Depends(get_current_user)):
    case = session.get(Case, case_id)
    if not case or case.org_id != user.org_id:
        raise HTTPException(404)

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    except ImportError:
        raise HTTPException(503, "reportlab not installed")

    # Generate AI-enriched report content
    ai = ai_enrich_case(case)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    r, g, b = severity_color_rgb(case.severity)
    accent = colors.Color(r/255, g/255, b/255)
    dark   = colors.Color(0.04, 0.03, 0.06)
    gray   = colors.Color(0.44, 0.44, 0.45)
    lgray  = colors.Color(0.97, 0.97, 0.98)

    def S(name, **kw):
        base = dict(fontName="Helvetica", fontSize=9, textColor=gray, spaceAfter=4, leading=14)
        base.update(kw)
        return ParagraphStyle(name, **base)

    title_style = S("title", fontName="Helvetica-Bold", fontSize=16, textColor=dark, spaceAfter=4)
    h2_style    = S("h2",    fontName="Helvetica-Bold", fontSize=11, textColor=accent, spaceBefore=14, spaceAfter=4)
    body_style  = S("body")
    mono_style  = S("mono",  fontName="Courier", fontSize=8, leading=12)
    note_style  = S("note",  fontName="Helvetica-Oblique", fontSize=8, textColor=gray)

    def clean(text):
        return (text or "").replace("\n", "<br/>").replace("**", "").replace("##", "").replace("#", "")

    story = []

    # ── Cover ─────────────────────────────────────────────────────────────────
    story.append(Paragraph("AEGISTRACE SOC PLATFORM", S("brand", fontName="Helvetica-Bold", fontSize=8, textColor=accent, spaceAfter=2)))
    story.append(Paragraph("INCIDENT INVESTIGATION REPORT", S("rep", fontName="Helvetica-Bold", fontSize=10, textColor=dark, spaceAfter=6)))
    story.append(Paragraph(case.title, title_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=accent, spaceAfter=10))

    detail_data = [
        ["Case Number",     case.case_number,    "Severity",        case.severity.upper()],
        ["Status",          case.status.replace("_"," ").upper(), "Incident Type", case.incident_type.replace("_"," ").title()],
        ["Analyst",         case.analyst_name,   "Customer",        case.customer_name or "—"],
        ["Classification",  case.classification, "Affected Systems",case.affected_systems or "—"],
        ["Report Date",     datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"), "Case Opened", case.created_at.strftime("%Y-%m-%d %H:%M UTC")],
    ]
    dt = Table(detail_data, colWidths=[3.5*cm, 5.5*cm, 3.5*cm, 5*cm])
    dt.setStyle(TableStyle([
        ("FONTNAME",  (0,0), (0,-1), "Helvetica-Bold"),
        ("FONTNAME",  (2,0), (2,-1), "Helvetica-Bold"),
        ("FONTNAME",  (1,0), (1,-1), "Helvetica"),
        ("FONTNAME",  (3,0), (3,-1), "Helvetica"),
        ("FONTSIZE",  (0,0), (-1,-1), 8),
        ("ROWBACKGROUNDS", (0,0), (-1,-1), [lgray, colors.white]),
        ("GRID",      (0,0), (-1,-1), 0.5, colors.Color(0.85,0.85,0.86)),
        ("PADDING",   (0,0), (-1,-1), 6),
    ]))
    story.append(dt)
    story.append(Spacer(1, 0.4*cm))

    # ── AI Executive Summary ──────────────────────────────────────────────────
    if ai.get("executive_summary"):
        story.append(Paragraph("Executive Summary", h2_style))
        story.append(Paragraph(clean(ai["executive_summary"]), body_style))

    # ── Risk Rating ───────────────────────────────────────────────────────────
    if ai.get("risk_rating"):
        story.append(Paragraph("Risk Rating", h2_style))
        story.append(Paragraph(clean(ai["risk_rating"]), body_style))
        if case.ai_severity_score:
            story.append(Paragraph(f"AI Severity Score: {case.ai_severity_score}/100  |  {case.ai_severity_reasoning or ''}", note_style))

    # ── Attack Timeline Summary ───────────────────────────────────────────────
    if ai.get("attack_timeline_summary"):
        story.append(Paragraph("Attack Timeline Summary", h2_style))
        story.append(Paragraph(clean(ai["attack_timeline_summary"]), body_style))

    # ── Technical Narrative ───────────────────────────────────────────────────
    if ai.get("technical_narrative"):
        story.append(Paragraph("Technical Analysis", h2_style))
        story.append(Paragraph(clean(ai["technical_narrative"]), body_style))

    # ── Threat Actor Profile ──────────────────────────────────────────────────
    if ai.get("threat_actor_profile"):
        story.append(Paragraph("Threat Actor Profile", h2_style))
        story.append(Paragraph(clean(ai["threat_actor_profile"]), body_style))

    # ── Impact Assessment ─────────────────────────────────────────────────────
    if ai.get("impact_assessment"):
        story.append(Paragraph("Impact Assessment", h2_style))
        story.append(Paragraph(clean(ai["impact_assessment"]), body_style))

    # ── Analyst Findings (raw) ────────────────────────────────────────────────
    if case.findings:
        story.append(Paragraph("Analyst Findings", h2_style))
        story.append(Paragraph(clean(case.findings), body_style))

    # IOCs
    try:
        iocs = json.loads(case.iocs or "[]")
        if iocs:
            story.append(Paragraph("Indicators of Compromise", h2_style))
            ioc_data = [["IOC", "Type", "Verdict", "Confidence"]]
            for ioc in iocs[:20]:
                ioc_data.append([
                    ioc.get("ioc", "")[:50],
                    ioc.get("type", ""),
                    ioc.get("verdict", ""),
                    f"{ioc.get('confidence', 0)}%",
                ])
            ti = Table(ioc_data, colWidths=[9*cm, 2*cm, 3*cm, 3*cm])
            ti.setStyle(TableStyle([
                ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
                ("FONTSIZE", (0,0), (-1,-1), 8),
                ("BACKGROUND", (0,0), (-1,0), colors.Color(0.07, 0.03, 0.06)),
                ("TEXTCOLOR", (0,0), (-1,0), colors.white),
                ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.Color(0.97,0.97,0.98), colors.white]),
                ("GRID", (0,0), (-1,-1), 0.5, colors.Color(0.85, 0.85, 0.86)),
                ("PADDING", (0,0), (-1,-1), 5),
                ("FONTNAME", (0,1), (0,-1), "Courier"),
            ]))
            story.append(ti)
    except Exception:
        pass

    # MITRE
    try:
        mitre = json.loads(case.mitre_techniques or "[]")
        if mitre:
            story.append(Paragraph("MITRE ATT&CK Techniques", h2_style))
            mitre_data = [["Technique ID", "Name", "Tactic"]]
            for m in mitre:
                mitre_data.append([m.get("id", ""), m.get("name", ""), m.get("tactic", "")])
            tm = Table(mitre_data, colWidths=[3*cm, 8*cm, 6*cm])
            tm.setStyle(TableStyle([
                ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
                ("FONTSIZE", (0,0), (-1,-1), 8),
                ("BACKGROUND", (0,0), (-1,0), colors.Color(0.33, 0.29, 0.73)),
                ("TEXTCOLOR", (0,0), (-1,0), colors.white),
                ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.Color(0.97,0.97,0.98), colors.white]),
                ("GRID", (0,0), (-1,-1), 0.5, colors.Color(0.85, 0.85, 0.86)),
                ("PADDING", (0,0), (-1,-1), 5),
                ("FONTNAME", (0,1), (0,-1), "Courier"),
            ]))
            story.append(tm)
    except Exception:
        pass

    # ── Remediation ───────────────────────────────────────────────────────────
    if ai.get("remediation_summary"):
        story.append(Paragraph("Remediation Actions", h2_style))
        story.append(Paragraph(clean(ai["remediation_summary"]), body_style))
    elif case.recommendations:
        story.append(Paragraph("Recommendations", h2_style))
        story.append(Paragraph(clean(case.recommendations), body_style))

    # ── Lessons Learned ───────────────────────────────────────────────────────
    if ai.get("lessons_learned"):
        story.append(Paragraph("Lessons Learned", h2_style))
        story.append(Paragraph(clean(ai["lessons_learned"]), body_style))

    # ── Closure Notes ─────────────────────────────────────────────────────────
    try:
        closure = json.loads(case.closure_notes or "{}")
        if closure:
            story.append(Paragraph("Case Closure", h2_style))
            if closure.get("final_findings"):
                story.append(Paragraph(f"<b>Final Findings:</b> {closure['final_findings']}", body_style))
            story.append(Paragraph(f"<b>Closed by:</b> {closure.get('closed_by','—')}  |  <b>Closed at:</b> {closure.get('closed_at','—')}", note_style))
    except Exception:
        pass

    # ── Commands Run ──────────────────────────────────────────────────────────
    if case.commands_run:
        story.append(Paragraph("Commands &amp; Tool Output", h2_style))
        for line in case.commands_run.split("\n")[:40]:
            story.append(Paragraph(line or "&nbsp;", mono_style))

    # ── Timeline ──────────────────────────────────────────────────────────────
    tl_events = session.exec(
        select(TimelineEvent)
        .where(TimelineEvent.case_id == case_id)
        .order_by(TimelineEvent.timestamp)
    ).all()
    if tl_events:
        story.append(Paragraph("Investigation Timeline", h2_style))
        tl_data = [["Timestamp", "Type", "Description"]]
        for ev in tl_events:
            tl_data.append([
                ev.timestamp.strftime("%Y-%m-%d %H:%M"),
                ev.event_type.upper(),
                (ev.description or "")[:80],
            ])
        tlt = Table(tl_data, colWidths=[3.5*cm, 2.5*cm, 11.5*cm])
        tlt.setStyle(TableStyle([
            ("FONTNAME",  (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE",  (0,0), (-1,-1), 7.5),
            ("BACKGROUND",(0,0), (-1,0), dark),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [lgray, colors.white]),
            ("GRID",      (0,0), (-1,-1), 0.5, colors.Color(0.85,0.85,0.86)),
            ("PADDING",   (0,0), (-1,-1), 4),
            ("FONTNAME",  (0,1), (0,-1), "Courier"),
        ]))
        story.append(tlt)

    # ── Footer ────────────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.6*cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=gray))
    story.append(Paragraph(
        f"Generated by AegisTrace SOC Platform &nbsp;·&nbsp; "
        f"{datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')} &nbsp;·&nbsp; "
        f"AI-Assisted Analysis (Groq llama-3.3-70b) &nbsp;·&nbsp; CONFIDENTIAL",
        S("footer", fontName="Helvetica", fontSize=7, textColor=gray, alignment=TA_CENTER)
    ))

    doc.build(story)
    buffer.seek(0)
    filename = f"AegisTrace_Report_{case.case_number}.pdf"
    return StreamingResponse(buffer, media_type="application/pdf",
                             headers={"Content-Disposition": f'attachment; filename="{filename}"'})


@router.get("/{case_id}/docx")
def download_docx(case_id: int, session: Session = Depends(get_session),
                  user: User = Depends(get_current_user)):
    case = session.get(Case, case_id)
    if not case or case.org_id != user.org_id:
        raise HTTPException(404)
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(503, "python-docx not installed")

    doc = Document()
    r, g, b = severity_color_rgb(case.severity)

    # Title
    h = doc.add_heading(case.title, 0)
    h.runs[0].font.color.rgb = RGBColor(r, g, b)
    doc.add_paragraph(f"Case: {case.case_number}  |  Severity: {case.severity.upper()}  |  Status: {case.status.upper()}")
    doc.add_paragraph(f"Analyst: {case.analyst_name}  |  Customer: {case.customer_name}")
    doc.add_paragraph(f"Type: {case.incident_type}  |  Date: {case.created_at.strftime('%Y-%m-%d')}")
    doc.add_paragraph("")

    def add_section(title, content):
        if not content:
            return
        doc.add_heading(title, level=2)
        doc.add_paragraph(content)

    add_section("Description", case.description)
    add_section("Executive Summary", case.ai_executive_summary)
    add_section("Findings", case.findings)
    add_section("Recommendations", case.recommendations)

    try:
        iocs = json.loads(case.iocs or "[]")
        if iocs:
            doc.add_heading("Indicators of Compromise", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "IOC", "Type", "Verdict", "Confidence"
            for ioc in iocs[:20]:
                row = table.add_row().cells
                row[0].text = ioc.get("ioc", "")
                row[1].text = ioc.get("type", "")
                row[2].text = ioc.get("verdict", "")
                row[3].text = f"{ioc.get('confidence', 0)}%"
    except Exception:
        pass

    if case.commands_run:
        doc.add_heading("Commands Run", level=2)
        p = doc.add_paragraph(case.commands_run)
        p.runs[0].font.name = "Courier New"
        p.runs[0].font.size = Pt(9)

    doc.add_paragraph(f"\nGenerated by AegisTrace · {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')} · CONFIDENTIAL")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"aegistrace_{case.case_number}.docx"
    return StreamingResponse(buffer,
                             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": f'attachment; filename="{filename}"'})


# ── DORA Compliance Report ────────────────────────────────────────────────────
@router.get("/{case_id}/dora")
def download_dora(case_id: int, session: Session = Depends(get_session),
                  user: User = Depends(get_current_user)):
    """
    Generate a DORA Article 19 Major ICT-related Incident Report.
    Maps the 5 DORA pillars to AegisTrace case data.
    """
    case = session.get(Case, case_id)
    if not case or case.org_id != user.org_id:
        raise HTTPException(404)

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.enums import TA_CENTER
    except ImportError:
        raise HTTPException(503, "reportlab not installed")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=2.2*cm, leftMargin=2.2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    dark   = colors.Color(0.04, 0.03, 0.06)
    red    = colors.Color(0.75, 0.22, 0.17)
    gray   = colors.Color(0.44, 0.44, 0.45)
    light  = colors.Color(0.97, 0.97, 0.98)

    def S(name, **kw):
        base = {"fontName": "Helvetica", "fontSize": 9, "textColor": gray, "spaceAfter": 4, "leading": 14}
        base.update(kw)
        return ParagraphStyle(name, **base)

    h1   = S("h1", fontName="Helvetica-Bold", fontSize=14, textColor=dark, spaceAfter=6)
    h2   = S("h2", fontName="Helvetica-Bold", fontSize=11, textColor=red, spaceBefore=16, spaceAfter=4)
    body = S("body")
    mono = S("mono", fontName="Courier", fontSize=8, textColor=gray, leading=12)
    note = S("note", fontSize=8, textColor=gray, fontName="Helvetica-Oblique")

    # Parse data
    closure = {}
    try:
        closure = json.loads(case.closure_notes or "{}")
    except Exception:
        pass
    iocs = []
    try:
        iocs = json.loads(case.iocs or "[]")
    except Exception:
        pass
    mitre = []
    try:
        mitre = json.loads(case.mitre_techniques or "[]")
    except Exception:
        pass

    # Timeline from DB
    timeline = session.exec(
        select(TimelineEvent)
        .where(TimelineEvent.case_id == case_id)
        .order_by(TimelineEvent.timestamp)
    ).all()

    first_detected = timeline[0].timestamp.strftime("%Y-%m-%d %H:%M UTC") if timeline else case.created_at.strftime("%Y-%m-%d %H:%M UTC")
    last_event     = timeline[-1].timestamp.strftime("%Y-%m-%d %H:%M UTC") if timeline else "Ongoing"

    story = []

    # EU / DORA Header
    story.append(Paragraph("EUROPEAN UNION — DIGITAL OPERATIONAL RESILIENCE ACT (DORA)", S("eu", fontName="Helvetica-Bold", fontSize=8, textColor=gray, spaceAfter=2)))
    story.append(Paragraph("MAJOR ICT-RELATED INCIDENT REPORT — Article 19(1)", h1))
    story.append(HRFlowable(width="100%", thickness=1.5, color=red, spaceAfter=10))

    story.append(Paragraph("REPORT IDENTIFICATION", h2))
    id_data = [
        ["Report Reference", case.case_number],
        ["Incident Title",   case.title],
        ["Classification",   case.classification.upper()],
        ["Report Date",      datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")],
        ["Entity Name",      case.customer_name or "—"],
        ["Prepared By",      case.analyst_name or "—"],
    ]
    t = Table(id_data, colWidths=[5.5*cm, 11*cm])
    t.setStyle(TableStyle([
        ("FONTNAME", (0,0), (0,-1), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 8),
        ("ROWBACKGROUNDS", (0,0), (-1,-1), [light, colors.white]),
        ("GRID", (0,0), (-1,-1), 0.5, colors.Color(0.85,0.85,0.86)),
        ("PADDING", (0,0), (-1,-1), 5),
    ]))
    story.append(t)

    # Pillar 1 — ICT Risk Management
    story.append(Paragraph("PILLAR 1 — ICT RISK MANAGEMENT", h2))
    story.append(Paragraph(f"<b>Incident Type:</b> {case.incident_type.replace('_',' ').title()}", body))
    story.append(Paragraph(f"<b>Severity Classification:</b> {case.severity.upper()}", body))
    story.append(Paragraph(f"<b>Affected Systems / Services:</b> {case.affected_systems or 'Not specified'}", body))
    story.append(Paragraph(f"<b>Description:</b>", body))
    story.append(Paragraph(case.description or "—", body))

    # Pillar 2 — Incident Detection
    story.append(Paragraph("PILLAR 2 — INCIDENT DETECTION & CLASSIFICATION", h2))
    story.append(Paragraph(f"<b>Initial Detection:</b> {first_detected}", body))
    story.append(Paragraph(f"<b>Detection Method:</b> {timeline[0].description if timeline else 'Manual detection'}", body))
    story.append(Paragraph(f"<b>AI Severity Score:</b> {case.ai_severity_score or 'Not assessed'}/100", body))
    if case.ai_severity_reasoning:
        story.append(Paragraph(f"<b>Severity Reasoning:</b> {case.ai_severity_reasoning}", body))

    # Pillar 3 — Incident Reporting
    story.append(Paragraph("PILLAR 3 — INCIDENT REPORTING & NOTIFICATION", h2))
    story.append(Paragraph(f"<b>Case Opened:</b> {case.created_at.strftime('%Y-%m-%d %H:%M UTC')}", body))
    story.append(Paragraph(f"<b>Last Updated:</b> {case.updated_at.strftime('%Y-%m-%d %H:%M UTC')}", body))
    story.append(Paragraph(f"<b>Current Status:</b> {case.status.replace('_',' ').upper()}", body))

    if mitre:
        story.append(Paragraph("<b>MITRE ATT&amp;CK Techniques Identified:</b>", body))
        m_data = [["Technique ID", "Name", "Tactic"]] + [[m.get("id",""), m.get("name",""), m.get("tactic","")] for m in mitre]
        mt = Table(m_data, colWidths=[3*cm, 9*cm, 4.5*cm])
        mt.setStyle(TableStyle([
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,-1), 8),
            ("BACKGROUND", (0,0), (-1,0), dark),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [light, colors.white]),
            ("GRID", (0,0), (-1,-1), 0.5, colors.Color(0.85,0.85,0.86)),
            ("PADDING", (0,0), (-1,-1), 4),
            ("FONTNAME", (0,1), (0,-1), "Courier"),
        ]))
        story.append(mt)

    # Pillar 4 — Resilience Testing (IOC evidence)
    story.append(Paragraph("PILLAR 4 — DIGITAL OPERATIONAL RESILIENCE TESTING", h2))
    story.append(Paragraph(f"<b>Indicators of Compromise ({len(iocs)}):</b>", body))
    if iocs:
        ioc_data = [["IOC Value", "Type", "Verdict"]] + [
            [i.get("ioc","")[:55], i.get("type",""), i.get("verdict","—")] for i in iocs[:20]
        ]
        it = Table(ioc_data, colWidths=[9.5*cm, 2.5*cm, 4.5*cm])
        it.setStyle(TableStyle([
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,-1), 8),
            ("BACKGROUND", (0,0), (-1,0), dark),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [light, colors.white]),
            ("GRID", (0,0), (-1,-1), 0.5, colors.Color(0.85,0.85,0.86)),
            ("PADDING", (0,0), (-1,-1), 4),
            ("FONTNAME", (0,1), (0,-1), "Courier"),
        ]))
        story.append(it)

    # Pillar 5 — ICT Third-Party Risk / Recovery
    story.append(Paragraph("PILLAR 5 — RECOVERY & LESSONS LEARNED", h2))
    if closure:
        story.append(Paragraph(f"<b>Final Findings:</b>", body))
        story.append(Paragraph(closure.get("final_findings", "—"), body))
        story.append(Paragraph(f"<b>Remediation Steps Taken:</b>", body))
        story.append(Paragraph(closure.get("remediation_steps", "—"), body))
        story.append(Paragraph(f"<b>Lessons Learned:</b>", body))
        story.append(Paragraph(closure.get("lessons_learned", "—"), body))
        story.append(Paragraph(f"<b>Case Closed By:</b> {closure.get('closed_by','—')}  |  <b>Closed At:</b> {closure.get('closed_at','—')}", body))
    else:
        story.append(Paragraph("Case not yet closed. Closure notes pending.", note))

    if case.recommendations:
        story.append(Paragraph("<b>Recommendations:</b>", body))
        story.append(Paragraph(case.recommendations.replace("\n","<br/>").replace("**",""), body))

    # Executive Summary
    if case.ai_executive_summary:
        story.append(Paragraph("AI-ASSISTED EXECUTIVE SUMMARY", h2))
        story.append(Paragraph(case.ai_executive_summary, body))

    # Footer
    story.append(Spacer(1, 0.5*cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=gray))
    story.append(Paragraph(
        f"Generated by AegisTrace SOC Platform · {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')} · DORA Article 19 Format · CONFIDENTIAL",
        S("footer", fontName="Helvetica", fontSize=7, textColor=gray, alignment=TA_CENTER)
    ))

    doc.build(story)
    buffer.seek(0)
    filename = f"DORA_{case.case_number}.pdf"
    return StreamingResponse(buffer, media_type="application/pdf",
                             headers={"Content-Disposition": f'attachment; filename="{filename}"'})


# ── DPDPA 2023 Compliance Report ──────────────────────────────────────────────
@router.get("/dpdpa/{case_id}")
def generate_dpdpa_report(
    case_id: int,
    _user: User = Depends(get_current_user),
    session: Session = Depends(get_session),
):
    """Generate DPDPA 2023 compliance incident report for a case."""
    case = session.get(Case, case_id)
    if not case or case.org_id != _user.org_id:
        raise HTTPException(404, "Case not found")

    iocs = json.loads(case.iocs or "[]")
    mitre = json.loads(case.mitre_techniques or "[]")
    timeline = session.exec(
        select(TimelineEvent).where(TimelineEvent.case_id == case_id)
        .order_by(TimelineEvent.timestamp.asc())
    ).all()

    # Build DPDPA-mapped findings
    dpdpa_obligations = [
        {
            "section": "Section 8(6)",
            "title": "Personal Data Breach Notification",
            "obligation": "Data Fiduciary must notify the Data Protection Board and affected Data Principals of any personal data breach",
            "status": "APPLICABLE" if case.severity in ("critical", "high") else "REVIEW",
            "evidence": f"Case {case.case_number}: {case.title}. Severity: {case.severity.upper()}",
        },
        {
            "section": "Section 8(5)",
            "title": "Data Retention and Erasure",
            "obligation": "Erase personal data when the purpose for processing is no longer served",
            "status": "REVIEW",
            "evidence": f"Investigation identified {len(iocs)} indicators. Review data retention policies.",
        },
        {
            "section": "Section 5",
            "title": "Grounds for Processing Personal Data",
            "obligation": "Processing of personal data must be for a lawful purpose with consent or legitimate use",
            "status": "REVIEW",
            "evidence": f"MITRE techniques involved: {', '.join(t.get('id','') for t in mitre[:5] if isinstance(t,dict)) or 'None identified'}",
        },
        {
            "section": "Section 10",
            "title": "Additional Obligations of Significant Data Fiduciaries",
            "obligation": "Conduct periodic Data Protection Impact Assessments",
            "status": "PENDING",
            "evidence": "Post-incident DPIA recommended based on breach scope",
        },
        {
            "section": "Section 77",
            "title": "Penalty for Breach of Obligations",
            "obligation": "Penalties up to INR 250 crore for failure to notify breach",
            "status": "INFORMATION",
            "evidence": f"Breach detected {case.created_at.strftime('%Y-%m-%d')}. Notification timeline must be documented.",
        },
    ]

    prompt = f"""Generate an executive summary for a DPDPA 2023 incident report.

Case: {case.case_number} — {case.title}
Severity: {case.severity}
Description: {(case.description or '')[:500]}
Findings: {(case.findings or '')[:400]}
Timeline events: {len(timeline)}
IOCs: {len(iocs)}

Write 3 paragraphs:
1. Incident overview and personal data impact assessment
2. DPDPA notification obligations triggered and timelines
3. Recommended remediation steps to satisfy DPDPA requirements

Be specific, professional, and reference DPDPA 2023 sections where relevant."""

    summary = call_ai("report", prompt)

    report_data = {
        "report_type": "DPDPA_2023",
        "case_number": case.case_number,
        "case_title": case.title,
        "severity": case.severity,
        "generated_at": datetime.utcnow().isoformat(),
        "generated_by": _user.name or _user.email,
        "organisation": _user.org_id,
        "incident_date": case.created_at.isoformat() if case.created_at else None,
        "executive_summary": summary,
        "dpdpa_obligations": dpdpa_obligations,
        "ioc_count": len(iocs),
        "timeline_events": len(timeline),
        "mitre_techniques": [t.get("id") for t in mitre if isinstance(t, dict)],
        "notification_deadline": "72 hours from detection (Section 8(6))",
        "data_protection_board": "dpboard.gov.in",
    }

    session.add(AuditLog(
        action="dpdpa_report_generated",
        entity_type="case",
        entity_id=str(case.case_number),
        user_id=_user.id,
        user_email=_user.email,
    ))
    session.commit()
    return report_data


@router.get("/regulatory-package/{case_id}")
async def generate_regulatory_package(
    case_id: int,
    regulation: str = Query(default="all"),  # all | eu_ai_act | dora | dpdpa
    _user: User = Depends(get_current_user),
    session: Session = Depends(get_session),
):
    """
    One-click Regulatory Evidence Package.

    Bundles everything an AI Act / DORA / DPDPA auditor needs:
    - Hash-chained AI decision log with chain integrity proof
    - Human approval records
    - ITDR alerts during the incident window
    - Trust Certificate (chain fingerprint)
    - Regulatory article mapping
    - Evidence completeness score

    This is the "mathematical audit integrity" export described in the ATSP standard.
    """
    from models import (
        ProvenanceLedger, TrustEvent, ITDRAlert, TimelineEvent,
    )
    from routers.provenance import _compute_entry_hash

    case = session.get(Case, case_id)
    if not case or case.org_id != _user.org_id:
        raise HTTPException(404, "Case not found")

    # ── 1. Provenance entries ────────────────────────────────────────────────
    provenance = session.exec(
        select(ProvenanceLedger)
        .where(ProvenanceLedger.case_id == case_id)
        .order_by(ProvenanceLedger.timestamp.asc())
    ).all()

    # ── 2. Chain verification ────────────────────────────────────────────────
    chain_valid  = True
    broken_at    = None
    legacy_count = 0
    for entry in provenance:
        if not entry.entry_hash:
            legacy_count += 1
            continue
        if _compute_entry_hash(entry) != entry.entry_hash:
            chain_valid = False
            broken_at   = entry.id
            break

    chain_fingerprint = hashlib.sha256(
        "|".join(e.entry_hash for e in provenance if e.entry_hash).encode()
    ).hexdigest()

    # ── 3. Human approval records ────────────────────────────────────────────
    approvals = [
        p for p in provenance
        if p.approval_status in ("approved", "rejected")
    ]

    # ── 4. AI decisions (auto-approved entries with model info) ──────────────
    ai_decisions = [
        p for p in provenance
        if p.model_used and p.action_type in ("ai_analysis", "ai_triage", "ai_summary")
    ]

    # ── 5. ITDR alerts during the case window ────────────────────────────────
    window_start = case.created_at - timedelta(hours=1) if case.created_at else datetime.utcnow() - timedelta(days=1)
    window_end   = datetime.utcnow()

    itdr_alerts = session.exec(
        select(ITDRAlert)
        .where(ITDRAlert.org_id == _user.org_id)
        .where(
            (ITDRAlert.case_id == case_id) |
            (ITDRAlert.detected_at.between(window_start, window_end))
        )
        .order_by(ITDRAlert.detected_at.asc())
    ).all()

    # ── 6. Timeline events ───────────────────────────────────────────────────
    timeline = session.exec(
        select(TimelineEvent)
        .where(TimelineEvent.case_id == case_id)
        .order_by(TimelineEvent.timestamp.asc())
    ).all()

    # ── 7. Delegation tokens covering this case's period ────────────────────
    delegation_tokens = []
    try:
        from models import AgentDelegationToken
        delegation_tokens = session.exec(
            select(AgentDelegationToken)
            .where(AgentDelegationToken.org_id == _user.org_id)
            .where(AgentDelegationToken.not_before <= window_end)
            .where(AgentDelegationToken.not_after >= window_start)
        ).all()
    except Exception:
        delegation_tokens = []  # table may not exist yet in older deployments

    # ── 8. Evidence completeness scoring ────────────────────────────────────
    checks = {
        "has_ai_decisions":      len(ai_decisions) > 0,
        "has_human_approvals":   len(approvals) > 0,
        "chain_intact":          chain_valid,
        "has_timeline":          len(timeline) > 0,
        "has_itdr_context":      len(itdr_alerts) > 0,
        "case_has_findings":     bool(case.findings),
        "case_has_mitre":        bool(case.mitre_techniques and case.mitre_techniques != "[]"),
        "has_delegation_tokens": len(delegation_tokens) > 0,
    }
    score = int(sum(checks.values()) / len(checks) * 100)

    # ── 9. Regulatory article mapping ───────────────────────────────────────
    EU_AI_ACT_ARTICLES = {
        "Article 9 — Risk Management":
            "chain_intact and has_ai_decisions",
        "Article 12 — Record Keeping":
            "has_ai_decisions and has_timeline",
        "Article 13 — Transparency":
            "has_ai_decisions",
        "Article 14 — Human Oversight":
            "has_human_approvals",
        "Article 17 — Quality Management":
            "has_delegation_tokens and chain_intact",
        "Article 26 — Obligations of Deployers":
            "has_delegation_tokens and has_human_approvals",
    }

    DORA_ARTICLES = {
        "Article 17 — ICT-related incident management":
            "has_timeline and has_itdr_context",
        "Article 19 — Major incident reporting":
            "has_findings and has_timeline",
        "Article 28 — General principles on ICT third-party risk":
            "has_delegation_tokens",
    }

    DPDPA_ARTICLES = {
        "Section 8(6) — Personal data breach notification":
            "has_timeline",
        "Section 10 — Data Protection Impact Assessment":
            "chain_intact and has_ai_decisions",
    }

    def _eval_mapping(mapping: dict) -> dict:
        result = {}
        # build eval namespace from checks, normalising key names
        ns = {k: v for k, v in checks.items()}
        ns["has_findings"] = checks.get("case_has_findings", False)
        for article, condition in mapping.items():
            try:
                satisfied = eval(condition, {"__builtins__": {}}, ns)
            except Exception:
                satisfied = False
            result[article] = {
                "satisfied": satisfied,
                "status": "Evidenced" if satisfied else "Gap",
            }
        return result

    reg_mappings = {}
    if regulation in ("all", "eu_ai_act"):
        reg_mappings["eu_ai_act"] = _eval_mapping(EU_AI_ACT_ARTICLES)
    if regulation in ("all", "dora"):
        reg_mappings["dora"] = _eval_mapping(DORA_ARTICLES)
    if regulation in ("all", "dpdpa"):
        reg_mappings["dpdpa"] = _eval_mapping(DPDPA_ARTICLES)

    # ── 10. Assemble package ─────────────────────────────────────────────────
    package = {
        "package_type":      "AegisTrace Regulatory Evidence Package",
        "package_version":   "1.0",
        "atsp_standard":     "https://github.com/Prasanna-27eng/AegisTrace/blob/main/ATSP_SPEC.md",
        "generated_at":      datetime.utcnow().isoformat(),
        "generated_by":      _user.email,
        "regulation_scope":  regulation,

        "case": {
            "id":             case.id,
            "case_number":    case.case_number,
            "title":          case.title,
            "severity":       case.severity,
            "status":         case.status,
            "incident_type":  case.incident_type,
            "created_at":     case.created_at.isoformat() if case.created_at else None,
            "analyst":        case.analyst_name,
            "mitre_techniques": json.loads(case.mitre_techniques or "[]"),
        },

        "chain_integrity": {
            "valid":             chain_valid,
            "broken_at_entry":   broken_at,
            "chain_fingerprint": chain_fingerprint,
            "entries_chained":   len([e for e in provenance if e.entry_hash]),
            "legacy_entries":    legacy_count,
            "algorithm":         "SHA-256 HMAC chain (ATSP §2.1)",
            "verification_note": "Chain fingerprint uniquely identifies this case's AI decision history. Tamper any entry and the fingerprint changes.",
        },

        "ai_decisions": [
            {
                "id":              p.id,
                "action_type":     p.action_type,
                "model_used":      p.model_used,
                "actor":           p.actor,
                "confidence":      p.confidence,
                "output_summary":  (p.output_summary or "")[:500],
                "approval_status": p.approval_status,
                "approved_by":     p.approved_by,
                "timestamp":       p.timestamp.isoformat() if p.timestamp else None,
                "entry_hash":      p.entry_hash,
                "prev_hash":       p.prev_hash,
            }
            for p in ai_decisions
        ],

        "human_approvals": [
            {
                "id":           p.id,
                "action_type":  p.action_type,
                "actor":        p.actor,
                "decision":     p.approval_status,
                "approved_by":  p.approved_by,
                "timestamp":    p.timestamp.isoformat() if p.timestamp else None,
                "entry_hash":   p.entry_hash,
            }
            for p in approvals
        ],

        "delegation_tokens": [
            {
                "token_id":      t.token_id,
                "agent_name":    t.agent_name,
                "authorized_by": t.authorized_by,
                "capabilities":  json.loads(t.capabilities or "[]"),
                "not_before":    t.not_before.isoformat(),
                "not_after":     t.not_after.isoformat(),
                "actions_taken": t.actions_taken,
                "signature":     t.signature[:32] + "..." if t.signature else None,
            }
            for t in delegation_tokens
        ],

        "itdr_context": [
            {
                "id":           a.id,
                "alert_type":   a.alert_type,
                "severity":     a.severity,
                "identity":     a.identity_label,
                "description":  a.description[:200],
                "detected_at":  a.detected_at.isoformat() if a.detected_at else None,
            }
            for a in itdr_alerts[:50]
        ],

        "incident_timeline": [
            {
                "id":          e.id,
                "event_type":  e.event_type,
                "description": e.description[:200],
                "timestamp":   e.timestamp.isoformat() if e.timestamp else None,
                "actor":       e.actor,
            }
            for e in timeline[:100]
        ],

        "regulatory_mapping": reg_mappings,

        "evidence_completeness": {
            "score_pct": score,
            "checks":    checks,
            "summary":   (
                f"Evidence package is {score}% complete. "
                + (f"{sum(1 for v in checks.values() if not v)} gap(s) identified." if score < 100 else "All evidence checks passed.")
            ),
        },

        "notification_deadlines": {
            "eu_ai_act":  "72 hours from detection (Article 73)",
            "dora":       "4 hours for major ICT incidents (Article 19)",
            "dpdpa":      "72 hours (Section 8(6))",
        },
    }

    _audit(session, "regulatory_package_generated", "case", str(case.id),
           _user.id, _user.email,
           f"Regulatory evidence package ({regulation}) for {case.case_number}")
    session.commit()


# ── Feature 8: Compliance Evidence Export PDF ──────────────────────────────────
COMPLIANCE_FRAMEWORKS = {
    "soc2": {
        "name": "SOC 2 Type II",
        "controls": [
            ("CC6.1", "Logical and Physical Access Controls"),
            ("CC6.7", "Access to Information Assets"),
            ("CC7.2", "Security Incident Management"),
            ("CC7.4", "Response to Security Incidents"),
            ("CC9.2", "Risk Management Process"),
        ],
    },
    "iso27001": {
        "name": "ISO/IEC 27001:2022",
        "controls": [
            ("A.5.25", "Assessment and decision on information security events"),
            ("A.5.26", "Response to information security incidents"),
            ("A.5.28", "Collection of evidence"),
            ("A.8.15", "Logging"),
            ("A.8.16", "Monitoring activities"),
        ],
    },
    "nist": {
        "name": "NIST Cybersecurity Framework 2.0",
        "controls": [
            ("ID.RA", "Risk Assessment"),
            ("DE.AE", "Anomalies and Events"),
            ("RS.AN", "Incident Analysis"),
            ("RS.MI", "Incident Mitigation"),
            ("RC.RP", "Recovery Planning"),
        ],
    },
    "dora": {
        "name": "DORA (EU) 2022/2554",
        "controls": [
            ("Art. 17", "ICT-related incident management process"),
            ("Art. 18", "Classification of ICT-related incidents"),
            ("Art. 19", "Reporting of major ICT-related incidents"),
            ("Art. 20", "Harmonisation of reporting content"),
            ("Art. 28", "General principles"),
        ],
    },
}


@router.post("/compliance-evidence")
def compliance_evidence_export(
    data: dict,
    user: User = Depends(get_current_user),
    session: Session = Depends(get_session),
):
    """
    Generate an audit-ready PDF compliance evidence package.
    Body: {"framework": "soc2|iso27001|nist|dora"}
    Returns a PDF file download.
    """
    framework_id = (data.get("framework") or "soc2").lower()
    framework = COMPLIANCE_FRAMEWORKS.get(framework_id)
    if not framework:
        raise HTTPException(400, f"Unknown framework '{framework_id}'. Options: {list(COMPLIANCE_FRAMEWORKS.keys())}")

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError:
        raise HTTPException(503, "reportlab not installed")

    # Fetch all closed cases in last 90 days
    cutoff = datetime.utcnow() - timedelta(days=90)
    cases = session.exec(
        select(Case)
        .where(Case.org_id == user.org_id)
        .where(Case.created_at >= cutoff)
        .order_by(Case.created_at.desc())
    ).all()

    closed_cases = [c for c in cases if c.status == "closed"]
    open_cases   = [c for c in cases if c.status != "closed"]
    critical_c   = [c for c in cases if c.severity == "critical"]

    # Basic SLA numbers
    ttc_hours = []
    for c in closed_cases:
        if c.created_at and c.updated_at:
            h = (c.updated_at - c.created_at).total_seconds() / 3600
            if h > 0:
                ttc_hours.append(h)
    avg_ttc = round(sum(ttc_hours) / len(ttc_hours), 1) if ttc_hours else None

    # Audit log count
    audit_count = session.exec(
        select(AuditLog).where(AuditLog.timestamp >= cutoff)
    ).all()

    # Build PDF
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    dark  = colors.Color(0.06, 0.06, 0.10)
    blue  = colors.Color(0.29, 0.49, 0.78)
    gray  = colors.Color(0.45, 0.45, 0.45)
    light = colors.Color(0.96, 0.97, 0.99)
    green = colors.Color(0.13, 0.77, 0.37)
    red   = colors.Color(0.94, 0.27, 0.27)

    def sty(name, **kw):
        base = {"fontName": "Helvetica", "fontSize": 9, "textColor": gray, "spaceAfter": 4, "leading": 14}
        base.update(kw)
        return ParagraphStyle(name, **base)

    H1   = sty("H1", fontName="Helvetica-Bold", fontSize=16, textColor=dark, spaceAfter=6)
    H2   = sty("H2", fontName="Helvetica-Bold", fontSize=12, textColor=blue, spaceBefore=14, spaceAfter=4)
    H3   = sty("H3", fontName="Helvetica-Bold", fontSize=10, textColor=dark, spaceAfter=3)
    body = sty("body")
    mono = sty("mono", fontName="Courier", fontSize=8, leading=12)
    note = sty("note", fontSize=8, textColor=gray, fontName="Helvetica-Oblique")

    story = []
    W = 17 * cm  # content width

    def hr(): return HRFlowable(width="100%", thickness=0.5, color=colors.Color(0.85, 0.85, 0.90))

    # Cover
    story += [
        Paragraph("AegisTrace ITDR Platform", sty("cover_sub", fontSize=10, textColor=blue)),
        Paragraph(f"{framework['name']} Compliance Evidence Package", H1),
        Spacer(1, 0.3*cm),
        hr(),
        Spacer(1, 0.2*cm),
        Paragraph(f"<b>Organization:</b> {user.org_id} | <b>Generated:</b> {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')} | <b>Period:</b> Last 90 days", body),
        Paragraph(f"<b>Prepared by:</b> {user.name or user.email} | <b>Framework Version:</b> {framework['name']}", body),
        Spacer(1, 0.5*cm),
    ]

    # Executive Summary
    story += [
        Paragraph("Executive Summary", H2),
        Paragraph(
            f"This evidence package documents AegisTrace security incident management activities "
            f"for the 90-day period ending {datetime.utcnow().strftime('%Y-%m-%d')}. "
            f"During this period, {len(cases)} security cases were managed: "
            f"{len(closed_cases)} closed, {len(open_cases)} currently active, "
            f"{len(critical_c)} classified as critical severity.",
            body,
        ),
        Spacer(1, 0.3*cm),
    ]

    # Key Metrics Table
    metrics = [
        ["Metric", "Value"],
        ["Total Cases (90 days)", str(len(cases))],
        ["Cases Closed",          str(len(closed_cases))],
        ["Cases Open",            str(len(open_cases))],
        ["Critical Incidents",    str(len(critical_c))],
        ["Avg Close Time",        f"{avg_ttc}h" if avg_ttc else "—"],
        ["Audit Log Entries",     str(len(audit_count))],
    ]
    t = Table(metrics, colWidths=[8*cm, 5*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND",  (0, 0), (-1, 0), blue),
        ("TEXTCOLOR",   (0, 0), (-1, 0), colors.white),
        ("FONTNAME",    (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",    (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light]),
        ("GRID",        (0, 0), (-1, -1), 0.5, colors.Color(0.85, 0.85, 0.9)),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING",(0, 0), (-1, -1), 8),
        ("TOPPADDING",  (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",(0, 0),(-1, -1), 5),
    ]))
    story += [t, Spacer(1, 0.4*cm)]

    # Framework Control Mapping
    story += [Paragraph(f"{framework['name']} Control Mapping", H2)]
    for ctrl_id, ctrl_name in framework["controls"]:
        story += [
            Paragraph(f"<b>{ctrl_id}</b> — {ctrl_name}", H3),
            Paragraph(
                f"AegisTrace maintains automated evidence for this control: {len(audit_count)} "
                f"audit log entries capturing all access, incident, and response activities "
                f"within the review period. See case summary section for incident-level evidence.",
                body,
            ),
            Spacer(1, 0.15*cm),
        ]

    # Case Evidence Table
    story += [hr(), Paragraph("Incident Case Evidence (Last 90 Days)", H2)]
    if not cases:
        story.append(Paragraph("No cases in the review period.", body))
    else:
        case_rows = [["Case #", "Title", "Severity", "Status", "Closed"]]
        for c in cases[:40]:
            closed_date = ""
            if c.status == "closed" and c.updated_at:
                closed_date = c.updated_at.strftime("%Y-%m-%d")
            case_rows.append([
                c.case_number or "",
                (c.title or "")[:52],
                (c.severity or "").upper(),
                c.status or "",
                closed_date,
            ])
        ct = Table(case_rows, colWidths=[2.2*cm, 7*cm, 2.2*cm, 2.2*cm, 2.5*cm])
        ct.setStyle(TableStyle([
            ("BACKGROUND",  (0, 0), (-1, 0), blue),
            ("TEXTCOLOR",   (0, 0), (-1, 0), colors.white),
            ("FONTNAME",    (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",    (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light]),
            ("GRID",        (0, 0), (-1, -1), 0.4, colors.Color(0.88, 0.88, 0.92)),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING",  (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING",(0, 0),(-1, -1), 4),
        ]))
        story += [ct, Spacer(1, 0.3*cm)]

    # Footer declaration
    story += [
        hr(), Spacer(1, 0.2*cm),
        Paragraph(
            f"This evidence package was auto-generated by AegisTrace v5.4 on "
            f"{datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}. "
            f"It should be reviewed by a qualified compliance officer before submission to auditors.",
            sty("footer", fontSize=8, textColor=gray, fontName="Helvetica-Oblique"),
        ),
    ]

    doc.build(story)
    buf.seek(0)

    _audit(session, "compliance_evidence_exported", "org", str(user.org_id),
           user.id, user.email, f"Framework: {framework_id}")
    session.commit()

    filename = f"aegistrace-{framework_id}-evidence-{datetime.utcnow().strftime('%Y%m%d')}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

    return package
